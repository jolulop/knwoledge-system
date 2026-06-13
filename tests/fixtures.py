"""Deterministic fixture builders for Phase 2 extraction tests.

These produce real, parseable files for each supported format without adding any
runtime dependency: the PDF is assembled by hand with a correct xref table, and the
DOCX is built with python-docx (already a declared extraction dependency).
"""
from __future__ import annotations

from pathlib import Path


def write_markdown(path: Path) -> Path:
    path.write_text(
        "# Title\n\n"
        "First paragraph of the document.\n\n"
        "## Section\n\n"
        "Second paragraph under a section.\n",
        encoding="utf-8",
    )
    return path


def write_html(path: Path) -> Path:
    path.write_text(
        "<html><head><style>.x{}</style>"
        "<script>alert(1)</script></head><body>"
        "<h1>Heading</h1><p>An HTML paragraph.</p>"
        "<ul><li>one</li><li>two</li></ul>"
        "</body></html>",
        encoding="utf-8",
    )
    return path


def write_csv(path: Path) -> Path:
    path.write_text("name,score\nalice,10\nbob,20\n", encoding="utf-8")
    return path


def write_docx(path: Path) -> Path:
    from docx import Document

    document = Document()
    document.add_heading("Doc Heading", level=1)
    document.add_paragraph("A docx paragraph of body text.")
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "h1"
    table.rows[0].cells[1].text = "h2"
    table.rows[1].cells[0].text = "v1"
    table.rows[1].cells[1].text = "v2"
    document.save(str(path))
    return path


def make_pdf_bytes(pages: list[str]) -> bytes:
    """Build a minimal multi-page PDF whose pages carry the given text via a Tj op."""
    objects: dict[int, bytes] = {}
    n = len(pages)
    kid_ids = [4 + 2 * i for i in range(n)]
    objects[1] = b"<< /Type /Catalog /Pages 2 0 R >>"
    kids = " ".join(f"{k} 0 R" for k in kid_ids)
    objects[2] = f"<< /Type /Pages /Kids [{kids}] /Count {n} >>".encode()
    objects[3] = b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>"
    for i, text in enumerate(pages):
        pid, cid = 4 + 2 * i, 5 + 2 * i
        esc = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
        stream = f"BT /F1 24 Tf 72 700 Td ({esc}) Tj ET".encode()
        objects[pid] = (
            f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            f"/Contents {cid} 0 R /Resources << /Font << /F1 3 0 R >> >> >>"
        ).encode()
        objects[cid] = b"<< /Length %d >>\nstream\n%s\nendstream" % (len(stream), stream)

    out = bytearray(b"%PDF-1.4\n")
    offsets: dict[int, int] = {}
    for num in sorted(objects):
        offsets[num] = len(out)
        out += f"{num} 0 obj\n".encode() + objects[num] + b"\nendobj\n"
    xref_pos = len(out)
    max_num = max(objects)
    out += f"xref\n0 {max_num + 1}\n".encode()
    out += b"0000000000 65535 f \n"
    for num in range(1, max_num + 1):
        out += f"{offsets[num]:010d} 00000 n \n".encode()
    out += (
        f"trailer\n<< /Size {max_num + 1} /Root 1 0 R >>\n"
        f"startxref\n{xref_pos}\n%%EOF"
    ).encode()
    return bytes(out)


def write_pdf(path: Path, pages: list[str]) -> Path:
    path.write_bytes(make_pdf_bytes(pages))
    return path
