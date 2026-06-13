#!/usr/bin/env python3
"""DOCX extractor (Phase 2): python-docx text and inline tables.

The document body is walked in true reading order (paragraphs and tables interleaved)
by iterating the underlying XML children, so headings, prose, and tables keep their
original sequence. Heading paragraphs are detected from their style name and mapped to
ATX levels. Tables are rendered inline as GitHub-flavored Markdown prose; structured
per-table CSV files are intentionally not produced for DOCX (Phase 2 Plan §8). DOCX is
not paginated, so ``page`` stays null.
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.workers.chunking import Element
from app.workers.extractors import Extraction, collapse_ws, gfm_table, pkg_version

_HEADING_STYLE = re.compile(r"^Heading\s+([1-6])$", re.IGNORECASE)


def _heading_level(style_name: str | None) -> int | None:
    if not style_name:
        return None
    if style_name.strip().lower() == "title":
        return 1
    match = _HEADING_STYLE.match(style_name.strip())
    return int(match.group(1)) if match else None


def _table_markdown(table: Table) -> str:
    rows = [[cell.text for cell in row.cells] for row in table.rows]
    if not rows:
        return ""
    return gfm_table(rows[0], rows[1:])


def extract(path: Path) -> Extraction:
    document = Document(str(path))
    body = document.element.body
    para_map = {p._p: p for p in document.paragraphs}
    table_map = {t._tbl: t for t in document.tables}

    elements: list[Element] = []
    for child in body.iterchildren():
        if child in para_map:
            paragraph: Paragraph = para_map[child]
            text = collapse_ws(paragraph.text)
            if not text:
                continue
            level = _heading_level(getattr(paragraph.style, "name", None))
            if level is not None:
                elements.append(Element(kind="heading", text=text, level=level))
            else:
                elements.append(Element(kind="prose", text=text))
        elif child in table_map:
            md = _table_markdown(table_map[child])
            if md:
                elements.append(Element(kind="prose", text=md))

    return Extraction(
        elements=elements, tool="python-docx", tool_version=pkg_version("python-docx")
    )
